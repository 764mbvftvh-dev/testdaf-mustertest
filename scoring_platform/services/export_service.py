from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pypinyin import pinyin, Style

from scoring_platform.config import (
    SECTION_LABELS, DIMENSION_LABELS, DIMENSION_DESCRIPTIONS,
    TASK_LABELS,
)
from scoring_platform.services.report_builder import list_graded_attempts, build_exam_summary
from scoring_platform.services.auth import list_all_students
from scoring_platform.services.speaking_store import list_all_speaking

RED = RGBColor(0xCC, 0x00, 0x00)


def _student_name(student_id: str) -> str:
    for s in list_all_students():
        if s["student_id"] == student_id:
            return s["name"]
    return student_id


def _set_cell(cell, text: str, tdn_label: str = "") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if tdn_label == "U3":
        run.font.color.rgb = RED
        run.bold = True


def export_word_all_students(sort: str = "name") -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("TestDaF 成绩总览", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    students = list_all_students()
    all_results = list_graded_attempts()
    all_speaking = list_all_speaking()

    students_with_data = {r["student_id"] for r in all_results if r["student_id"]}
    if sort == "name":
        students_with_data = sorted(students_with_data, key=lambda sid: "".join(p[0] for p in pinyin(_student_name(sid), style=Style.TONE3)))
    elif sort == "time":
        students_with_data = sorted(students_with_data, key=lambda sid: max((r["submitted_at"] for r in all_results if r["student_id"] == sid), default=""), reverse=True)
    else:
        students_with_data = sorted(students_with_data, key=lambda sid: _student_name(sid))

    if not students_with_data:
        doc.add_paragraph("暂无成绩数据。")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "学生姓名"
    hdr[1].text = "阅读 TDN"
    hdr[2].text = "听力 TDN"
    hdr[3].text = "写作 TDN"
    hdr[4].text = "口语 TDN"
    hdr[5].text = "最近考试"
    hdr[6].text = "学生 ID"

    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    for sid in students_with_data:
        name = _student_name(sid)
        row = table.add_row()
        cells = row.cells

        latest = {"reading": ("—", ""), "listening": ("—", ""), "writing": ("—", ""), "speaking": ("—", "")}
        last_exam = ""
        for r in all_results:
            if r["student_id"] != sid:
                continue
            sec = r["section"]
            if sec in latest and latest[sec][0] == "—":
                latest[sec] = (r["tdn_label"], r["tdn_label"])
            if r["exam_id"] and not last_exam:
                last_exam = r["submitted_at"][:10]

        spk = all_speaking.get(sid, {})
        if spk.get("overall_label"):
            latest["speaking"] = (spk["overall_label"], spk["overall_label"])

        _set_cell(cells[0], name)
        _set_cell(cells[1], latest["reading"][0], latest["reading"][1])
        _set_cell(cells[2], latest["listening"][0], latest["listening"][1])
        _set_cell(cells[3], latest["writing"][0], latest["writing"][1])
        _set_cell(cells[4], latest["speaking"][0], latest["speaking"][1])
        _set_cell(cells[5], last_exam)
        _set_cell(cells[6], sid)

    doc.add_paragraph("")
    doc.add_paragraph(f"共 {len(students_with_data)} 名学生有成绩记录。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_word_student(student_id: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    name = _student_name(student_id)
    title = doc.add_heading(f"TestDaF 成绩报告 · {name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Student ID: {student_id}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    results = [r for r in list_graded_attempts(student_id) if r["section"] != "speaking"]

    for r in results:
        section = r["section"]
        icon, label = SECTION_LABELS.get(section, ("", section))
        doc.add_heading(f"{icon} {label}", level=2)
        doc.add_paragraph(f"题目: {r['title']}")
        doc.add_paragraph(f"时间: {r['submitted_at']}")

        tdn_para = doc.add_paragraph()
        tdn_run = tdn_para.add_run(f"评分: {r['tdn_label']}")
        if r["tdn_label"] == "U3":
            tdn_run.font.color.rgb = RED
            tdn_run.bold = True

        if section in ("reading", "listening") and r["correct"] is not None:
            doc.add_paragraph(f"正确: {r['correct']} / {r['total']}")
        if section == "writing" and r.get("dimensionen"):
            dim_table = doc.add_table(rows=1, cols=2)
            dim_table.style = "Table Grid"
            dh = dim_table.rows[0].cells
            dh[0].text = "评分维度"
            dh[1].text = "得分"
            for hc in dh:
                for p_c in hc.paragraphs:
                    for run in p_c.runs:
                        run.bold = True
            for key, dlabel in DIMENSION_LABELS.items():
                drow = dim_table.add_row()
                drow.cells[0].text = dlabel
                drow.cells[1].text = str(r.get("dimensionen", {}).get(key, "—"))
            if r.get("kommentar"):
                doc.add_paragraph(f"评语: {r['kommentar']}")

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
